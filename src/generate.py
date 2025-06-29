from typing import List, Dict
import random

def generate_questions(topics: List[str], question_types: List[str], num_questions: int = 10) -> List[Dict]:
    """
    Generate sample questions for each topic using templates.
    Returns a list of dicts: {question, type, topic, bloom_level}
    """
    bloom_levels = ['Remember', 'Understand', 'Apply', 'Analyze', 'Evaluate', 'Create']
    templates = {
        'MCQ': 'Which of the following best describes {topic}?',
        'Short Answer': 'Explain the concept of {topic}.',
        'Long Answer': 'Discuss {topic} in detail with examples.',
        'Case Study': 'Given a scenario related to {topic}, analyze and provide solutions.'
    }
    questions = []
    for topic in topics:
        for qtype in question_types:
            for _ in range(num_questions // (len(topics) * len(question_types)) + 1):
                bloom = random.choice(bloom_levels)
                q = templates[qtype].format(topic=topic)
                questions.append({
                    'question': q,
                    'type': qtype,
                    'topic': topic,
                    'bloom_level': bloom
                })
    return questions[:num_questions]

def generate_model_answer(question_dict):
    """Generate a sample/model answer for a question dict."""
    qtype = question_dict['type']
    topic = question_dict['topic']
    if qtype == 'MCQ':
        return f"The correct answer is the option that best describes {topic}."
    elif qtype == 'Short Answer':
        return f"A short answer should explain the main idea of {topic} in 2-3 sentences."
    elif qtype == 'Long Answer':
        return f"A detailed answer should cover all aspects of {topic}, include examples, and demonstrate understanding."
    elif qtype == 'Case Study':
        return f"Analyze the scenario, relate it to {topic}, and provide a reasoned solution."
    else:
        return "Model answer not available."

def assign_marks(question_dict):
    qtype = question_dict['type']
    if qtype == 'MCQ':
        return 1
    elif qtype == 'Short Answer':
        return 3
    elif qtype == 'Long Answer':
        return 8
    elif qtype == 'Case Study':
        return 10
    else:
        return 2

def format_export_text(questions):
    lines = []
    total_marks = 0
    for idx, q in enumerate(questions, 1):
        marks = assign_marks(q)
        total_marks += marks
        lines.append(f"Q{idx} ({q['type']}, {q['bloom_level']}, {q['topic']}, {marks} marks): {q['question']}")
        lines.append(f"Model Answer: {generate_model_answer(q)}\n")
    lines.append(f"Total Assigned Marks: {total_marks}")
    return '\n'.join(lines)

# DOCX export
def format_export_docx(questions, filename='Question_Paper.docx'):
    from docx import Document
    doc = Document()
    doc.add_heading('Generated Question Paper', 0)
    total_marks = 0
    for idx, q in enumerate(questions, 1):
        marks = assign_marks(q)
        total_marks += marks
        doc.add_paragraph(f"Q{idx} ({q['type']}, {q['bloom_level']}, {q['topic']}, {marks} marks): {q['question']}", style='List Number')
        doc.add_paragraph(f"Model Answer: {generate_model_answer(q)}")
    # Add total marks as bold text
    p = doc.add_paragraph()
    run = p.add_run(f"Total Assigned Marks: {total_marks}")
    run.bold = True
    doc.save(filename)
    return filename 