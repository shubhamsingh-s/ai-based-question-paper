import pandas as pd
import numpy as np
from typing import Dict, List, Optional
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches
import io
import base64

class ReportGenerator:
    def __init__(self):
        self.report_data = {}
        self.visualizations = {}
        
    def generate_comprehensive_report(self, 
                                    question_paper: Dict,
                                    analysis_data: Dict,
                                    predictions: List[Dict] = None,
                                    historical_data: Dict = None) -> Dict:
        """Generate a comprehensive report with all analysis and visualizations"""
        
        report = {
            'metadata': {
                'generated_at': datetime.now().isoformat(),
                'paper_title': question_paper.get('paper_info', {}).get('title', 'Generated Question Paper'),
                'total_questions': question_paper.get('paper_info', {}).get('total_questions', 0),
                'total_marks': question_paper.get('paper_info', {}).get('total_marks', 0),
                'duration': question_paper.get('paper_info', {}).get('duration', 180)
            },
            'paper_analysis': self._analyze_question_paper(question_paper),
            'topic_analysis': self._analyze_topic_distribution(question_paper),
            'cognitive_analysis': self._analyze_cognitive_levels(question_paper),
            'difficulty_analysis': self._analyze_difficulty_distribution(question_paper),
            'predictions': predictions or [],
            'historical_comparison': self._compare_with_historical_data(question_paper, historical_data),
            'recommendations': self._generate_recommendations(question_paper, analysis_data),
            'visualizations': self._create_visualizations(question_paper, analysis_data, predictions)
        }
        
        return report
    
    def _analyze_question_paper(self, question_paper: Dict) -> Dict:
        """Analyze the generated question paper"""
        questions = question_paper.get('questions', [])
        
        if not questions:
            return {}
        
        # Basic statistics
        total_questions = len(questions)
        total_marks = sum(q.get('marks', 0) for q in questions)
        avg_marks_per_question = total_marks / total_questions if total_questions > 0 else 0
        
        # Question type distribution
        type_distribution = {}
        for q in questions:
            qtype = q.get('type', 'Unknown')
            type_distribution[qtype] = type_distribution.get(qtype, 0) + 1
        
        # Marks distribution by type
        type_marks = {}
        for q in questions:
            qtype = q.get('type', 'Unknown')
            marks = q.get('marks', 0)
            if qtype not in type_marks:
                type_marks[qtype] = []
            type_marks[qtype].append(marks)
        
        # Calculate average marks per type
        avg_marks_by_type = {}
        for qtype, marks_list in type_marks.items():
            avg_marks_by_type[qtype] = np.mean(marks_list)
        
        return {
            'total_questions': total_questions,
            'total_marks': total_marks,
            'average_marks_per_question': avg_marks_per_question,
            'question_type_distribution': type_distribution,
            'average_marks_by_type': avg_marks_by_type,
            'marks_range': {
                'min': min(q.get('marks', 0) for q in questions),
                'max': max(q.get('marks', 0) for q in questions)
            },
            'paper_balance_score': self._calculate_balance_score(questions)
        }
    
    def _analyze_topic_distribution(self, question_paper: Dict) -> Dict:
        """Analyze topic distribution in the question paper"""
        questions = question_paper.get('questions', [])
        
        topic_distribution = {}
        topic_marks = {}
        
        for q in questions:
            topic = q.get('topic', 'Unknown')
            marks = q.get('marks', 0)
            
            # Count questions per topic
            topic_distribution[topic] = topic_distribution.get(topic, 0) + 1
            
            # Sum marks per topic
            topic_marks[topic] = topic_marks.get(topic, 0) + marks
        
        # Calculate topic weightage
        total_marks = sum(topic_marks.values())
        topic_weightage = {}
        for topic, marks in topic_marks.items():
            topic_weightage[topic] = (marks / total_marks) * 100 if total_marks > 0 else 0
        
        # Topic coverage analysis
        unique_topics = len(topic_distribution)
        coverage_analysis = {
            'total_unique_topics': unique_topics,
            'topic_coverage_efficiency': self._calculate_coverage_efficiency(topic_distribution),
            'topic_balance_score': self._calculate_topic_balance(topic_distribution)
        }
        
        return {
            'topic_distribution': topic_distribution,
            'topic_marks_distribution': topic_marks,
            'topic_weightage': topic_weightage,
            'coverage_analysis': coverage_analysis
        }
    
    def _analyze_cognitive_levels(self, question_paper: Dict) -> Dict:
        """Analyze cognitive levels (Bloom's taxonomy) distribution"""
        questions = question_paper.get('questions', [])
        
        bloom_distribution = {}
        bloom_marks = {}
        
        for q in questions:
            bloom_level = q.get('bloom_level', 'Unknown')
            marks = q.get('marks', 0)
            
            bloom_distribution[bloom_level] = bloom_distribution.get(bloom_level, 0) + 1
            bloom_marks[bloom_level] = bloom_marks.get(bloom_level, 0) + marks
        
        # Calculate cognitive level weightage
        total_marks = sum(bloom_marks.values())
        bloom_weightage = {}
        for level, marks in bloom_marks.items():
            bloom_weightage[level] = (marks / total_marks) * 100 if total_marks > 0 else 0
        
        # Cognitive complexity analysis
        complexity_analysis = self._analyze_cognitive_complexity(bloom_distribution)
        
        return {
            'bloom_distribution': bloom_distribution,
            'bloom_marks_distribution': bloom_marks,
            'bloom_weightage': bloom_weightage,
            'complexity_analysis': complexity_analysis
        }
    
    def _analyze_difficulty_distribution(self, question_paper: Dict) -> Dict:
        """Analyze difficulty level distribution"""
        questions = question_paper.get('questions', [])
        
        difficulty_distribution = {}
        difficulty_marks = {}
        
        for q in questions:
            difficulty = q.get('difficulty', 'Medium')
            marks = q.get('marks', 0)
            
            difficulty_distribution[difficulty] = difficulty_distribution.get(difficulty, 0) + 1
            difficulty_marks[difficulty] = difficulty_marks.get(difficulty, 0) + marks
        
        # Calculate difficulty weightage
        total_marks = sum(difficulty_marks.values())
        difficulty_weightage = {}
        for level, marks in difficulty_marks.items():
            difficulty_weightage[level] = (marks / total_marks) * 100 if total_marks > 0 else 0
        
        # Difficulty balance analysis
        balance_analysis = self._analyze_difficulty_balance(difficulty_distribution)
        
        return {
            'difficulty_distribution': difficulty_distribution,
            'difficulty_marks_distribution': difficulty_marks,
            'difficulty_weightage': difficulty_weightage,
            'balance_analysis': balance_analysis
        }
    
    def _compare_with_historical_data(self, question_paper: Dict, historical_data: Dict) -> Dict:
        """Compare current paper with historical patterns"""
        if not historical_data:
            return {'status': 'No historical data available for comparison'}
        
        current_analysis = self._analyze_question_paper(question_paper)
        current_topic_analysis = self._analyze_topic_distribution(question_paper)
        
        comparison = {
            'topic_coverage_comparison': {},
            'type_distribution_comparison': {},
            'difficulty_comparison': {},
            'overall_similarity_score': 0
        }
        
        # Compare topic coverage
        if 'topic_weightage' in historical_data:
            current_topics = current_topic_analysis.get('topic_weightage', {})
            historical_topics = historical_data.get('topic_weightage', {})
            
            for topic in set(current_topics.keys()) | set(historical_topics.keys()):
                current_weight = current_topics.get(topic, 0)
                historical_weight = historical_topics.get(topic, 0)
                difference = current_weight - historical_weight
                
                comparison['topic_coverage_comparison'][topic] = {
                    'current_weightage': current_weight,
                    'historical_weightage': historical_weight,
                    'difference': difference,
                    'trend': 'increased' if difference > 0 else 'decreased' if difference < 0 else 'stable'
                }
        
        # Calculate overall similarity score
        similarity_score = self._calculate_similarity_score(current_analysis, historical_data)
        comparison['overall_similarity_score'] = similarity_score
        
        return comparison
    
    def _generate_recommendations(self, question_paper: Dict, analysis_data: Dict) -> List[Dict]:
        """Generate recommendations for improving the question paper"""
        recommendations = []
        
        # Analyze topic coverage
        topic_analysis = self._analyze_topic_distribution(question_paper)
        topic_weightage = topic_analysis.get('topic_weightage', {})
        
        if topic_weightage:
            # Check for over-represented topics
            for topic, weight in topic_weightage.items():
                if weight > 20:  # More than 20% weightage
                    recommendations.append({
                        'type': 'topic_balance',
                        'issue': f'Topic "{topic}" has high weightage ({weight:.1f}%)',
                        'recommendation': f'Consider reducing questions on {topic} to improve balance',
                        'priority': 'high' if weight > 30 else 'medium'
                    })
            
            # Check for under-represented topics
            low_weight_topics = [topic for topic, weight in topic_weightage.items() if weight < 5]
            if low_weight_topics:
                recommendations.append({
                    'type': 'topic_coverage',
                    'issue': f'Topics with low coverage: {", ".join(low_weight_topics)}',
                    'recommendation': 'Consider adding more questions on these topics',
                    'priority': 'medium'
                })
        
        # Analyze cognitive levels
        cognitive_analysis = self._analyze_cognitive_levels(question_paper)
        bloom_weightage = cognitive_analysis.get('bloom_weightage', {})
        
        # Check for cognitive level balance
        higher_order_skills = ['Analyze', 'Evaluate', 'Create']
        higher_order_weight = sum(bloom_weightage.get(level, 0) for level in higher_order_skills)
        
        if higher_order_weight < 30:
            recommendations.append({
                'type': 'cognitive_balance',
                'issue': f'Low coverage of higher-order thinking skills ({higher_order_weight:.1f}%)',
                'recommendation': 'Increase questions requiring analysis, evaluation, and creation',
                'priority': 'high'
            })
        
        # Analyze difficulty distribution
        difficulty_analysis = self._analyze_difficulty_distribution(question_paper)
        difficulty_weightage = difficulty_analysis.get('difficulty_weightage', {})
        
        hard_weight = difficulty_weightage.get('Hard', 0)
        if hard_weight < 10:
            recommendations.append({
                'type': 'difficulty_balance',
                'issue': f'Low coverage of difficult questions ({hard_weight:.1f}%)',
                'recommendation': 'Consider adding more challenging questions to test advanced understanding',
                'priority': 'medium'
            })
        
        return recommendations
    
    def _create_visualizations(self, question_paper: Dict, analysis_data: Dict, predictions: List[Dict]) -> Dict:
        """Create comprehensive visualizations for the report"""
        visualizations = {}
        
        # Topic distribution pie chart
        topic_analysis = self._analyze_topic_distribution(question_paper)
        topic_weightage = topic_analysis.get('topic_weightage', {})
        
        if topic_weightage:
            fig1 = px.pie(
                values=list(topic_weightage.values()),
                names=list(topic_weightage.keys()),
                title='Topic Weightage Distribution',
                hole=0.3
            )
            fig1.update_traces(textposition='inside', textinfo='percent+label')
            visualizations['topic_distribution'] = fig1
        
        # Question type distribution bar chart
        paper_analysis = self._analyze_question_paper(question_paper)
        type_distribution = paper_analysis.get('question_type_distribution', {})
        
        if type_distribution:
            fig2 = px.bar(
                x=list(type_distribution.keys()),
                y=list(type_distribution.values()),
                title='Question Type Distribution',
                color=list(type_distribution.values()),
                color_continuous_scale='viridis'
            )
            visualizations['type_distribution'] = fig2
        
        # Cognitive levels radar chart
        cognitive_analysis = self._analyze_cognitive_levels(question_paper)
        bloom_weightage = cognitive_analysis.get('bloom_weightage', {})
        
        if bloom_weightage:
            fig3 = go.Figure()
            fig3.add_trace(go.Scatterpolar(
                r=list(bloom_weightage.values()),
                theta=list(bloom_weightage.keys()),
                fill='toself',
                name='Cognitive Levels'
            ))
            fig3.update_layout(
                polar=dict(radialaxis=dict(visible=True, range=[0, max(bloom_weightage.values())])),
                showlegend=False,
                title='Cognitive Levels Distribution (Radar Chart)'
            )
            visualizations['cognitive_levels'] = fig3
        
        # Difficulty distribution
        difficulty_analysis = self._analyze_difficulty_distribution(question_paper)
        difficulty_weightage = difficulty_analysis.get('difficulty_weightage', {})
        
        if difficulty_weightage:
            fig4 = px.bar(
                x=list(difficulty_weightage.keys()),
                y=list(difficulty_weightage.values()),
                title='Difficulty Level Distribution',
                color=list(difficulty_weightage.values()),
                color_continuous_scale='RdYlGn'
            )
            visualizations['difficulty_distribution'] = fig4
        
        # Predictions visualization
        if predictions:
            pred_df = pd.DataFrame(predictions)
            fig5 = px.bar(
                data_frame=pred_df,
                x='topic',
                y='probability',
                title='Predicted Topic Probabilities',
                color='confidence',
                color_continuous_scale='plasma'
            )
            visualizations['predictions'] = fig5
        
        # Comprehensive dashboard
        if visualizations:
            dashboard = self._create_dashboard(visualizations)
            visualizations['dashboard'] = dashboard
        
        return visualizations
    
    def _create_dashboard(self, visualizations: Dict) -> go.Figure:
        """Create a comprehensive dashboard with multiple charts"""
        # Create subplots
        fig = make_subplots(
            rows=2, cols=2,
            subplot_titles=('Topic Distribution', 'Question Types', 'Cognitive Levels', 'Difficulty Levels'),
            specs=[[{"type": "pie"}, {"type": "bar"}],
                   [{"type": "bar"}, {"type": "bar"}]]
        )
        
        # Add traces from existing visualizations
        if 'topic_distribution' in visualizations:
            fig.add_trace(visualizations['topic_distribution'].data[0], row=1, col=1)
        
        if 'type_distribution' in visualizations:
            fig.add_trace(visualizations['type_distribution'].data[0], row=1, col=2)
        
        if 'cognitive_levels' in visualizations:
            fig.add_trace(visualizations['cognitive_levels'].data[0], row=2, col=1)
        
        if 'difficulty_distribution' in visualizations:
            fig.add_trace(visualizations['difficulty_distribution'].data[0], row=2, col=2)
        
        fig.update_layout(height=800, title_text="Question Paper Analysis Dashboard")
        return fig
    
    def export_report_to_docx(self, report: Dict, filename: str = 'question_paper_report.docx') -> str:
        """Export the report to a Word document"""
        doc = Document()
        
        # Title
        doc.add_heading('Question Paper Analysis Report', 0)
        
        # Metadata
        doc.add_heading('Report Information', level=1)
        metadata = report.get('metadata', {})
        doc.add_paragraph(f"Generated: {metadata.get('generated_at', 'Unknown')}")
        doc.add_paragraph(f"Paper Title: {metadata.get('paper_title', 'Unknown')}")
        doc.add_paragraph(f"Total Questions: {metadata.get('total_questions', 0)}")
        doc.add_paragraph(f"Total Marks: {metadata.get('total_marks', 0)}")
        doc.add_paragraph(f"Duration: {metadata.get('duration', 0)} minutes")
        
        # Paper Analysis
        doc.add_heading('Paper Analysis', level=1)
        paper_analysis = report.get('paper_analysis', {})
        doc.add_paragraph(f"Average marks per question: {paper_analysis.get('average_marks_per_question', 0):.2f}")
        doc.add_paragraph(f"Balance score: {paper_analysis.get('paper_balance_score', 0):.2f}")
        
        # Topic Analysis
        doc.add_heading('Topic Analysis', level=1)
        topic_analysis = report.get('topic_analysis', {})
        topic_weightage = topic_analysis.get('topic_weightage', {})
        
        for topic, weight in topic_weightage.items():
            doc.add_paragraph(f"{topic}: {weight:.1f}%")
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        recommendations = report.get('recommendations', [])
        
        for rec in recommendations:
            p = doc.add_paragraph()
            p.add_run(f"{rec['type'].title()}: ").bold = True
            p.add_run(rec['issue'])
            doc.add_paragraph(f"Recommendation: {rec['recommendation']}")
            doc.add_paragraph(f"Priority: {rec['priority']}")
        
        doc.save(filename)
        return filename
    
    def export_report_to_json(self, report: Dict, filename: str = 'question_paper_report.json') -> str:
        """Export the report to JSON format"""
        # Remove plotly figures as they can't be serialized
        export_report = report.copy()
        if 'visualizations' in export_report:
            export_report['visualizations'] = {
                'available_charts': list(report['visualizations'].keys())
            }
        
        with open(filename, 'w') as f:
            json.dump(export_report, f, indent=2, default=str)
        
        return filename
    
    # Helper methods
    def _calculate_balance_score(self, questions: List[Dict]) -> float:
        """Calculate balance score for the question paper"""
        if not questions:
            return 0.0
        
        # Calculate balance based on topic distribution
        topic_counts = {}
        for q in questions:
            topic = q.get('topic', 'Unknown')
            topic_counts[topic] = topic_counts.get(topic, 0) + 1
        
        if len(topic_counts) <= 1:
            return 0.0
        
        # Calculate coefficient of variation (lower is more balanced)
        counts = list(topic_counts.values())
        mean_count = np.mean(counts)
        std_count = np.std(counts)
        
        if mean_count == 0:
            return 0.0
        
        cv = std_count / mean_count
        # Convert to balance score (0-1, higher is more balanced)
        balance_score = max(0, 1 - cv)
        
        return balance_score
    
    def _calculate_coverage_efficiency(self, topic_distribution: Dict) -> float:
        """Calculate topic coverage efficiency"""
        if not topic_distribution:
            return 0.0
        
        total_questions = sum(topic_distribution.values())
        unique_topics = len(topic_distribution)
        
        # Efficiency = unique topics / total questions (normalized)
        efficiency = unique_topics / total_questions if total_questions > 0 else 0
        return min(1.0, efficiency * 2)  # Normalize to 0-1 range
    
    def _calculate_topic_balance(self, topic_distribution: Dict) -> float:
        """Calculate topic balance score"""
        if not topic_distribution:
            return 0.0
        
        counts = list(topic_distribution.values())
        if len(counts) <= 1:
            return 1.0  # Perfect balance if only one topic
        
        # Calculate Gini coefficient (lower is more balanced)
        sorted_counts = sorted(counts)
        n = len(sorted_counts)
        cumsum = np.cumsum(sorted_counts)
        gini = (n + 1 - 2 * np.sum(cumsum) / cumsum[-1]) / n if cumsum[-1] > 0 else 0
        
        # Convert to balance score (0-1, higher is more balanced)
        balance_score = 1 - gini
        return balance_score
    
    def _analyze_cognitive_complexity(self, bloom_distribution: Dict) -> Dict:
        """Analyze cognitive complexity of the question paper"""
        complexity_weights = {
            'Remember': 1,
            'Understand': 2,
            'Apply': 3,
            'Analyze': 4,
            'Evaluate': 5,
            'Create': 6
        }
        
        total_questions = sum(bloom_distribution.values())
        if total_questions == 0:
            return {'average_complexity': 0, 'complexity_level': 'Low'}
        
        weighted_sum = sum(bloom_distribution.get(level, 0) * complexity_weights.get(level, 1) 
                          for level in bloom_distribution)
        average_complexity = weighted_sum / total_questions
        
        # Determine complexity level
        if average_complexity <= 2:
            complexity_level = 'Low'
        elif average_complexity <= 4:
            complexity_level = 'Medium'
        else:
            complexity_level = 'High'
        
        return {
            'average_complexity': average_complexity,
            'complexity_level': complexity_level
        }
    
    def _analyze_difficulty_balance(self, difficulty_distribution: Dict) -> Dict:
        """Analyze difficulty balance"""
        if not difficulty_distribution:
            return {'balance_score': 0, 'recommendation': 'No difficulty data available'}
        
        # Ideal distribution: 30% Easy, 50% Medium, 20% Hard
        ideal_distribution = {'Easy': 0.3, 'Medium': 0.5, 'Hard': 0.2}
        
        total_questions = sum(difficulty_distribution.values())
        if total_questions == 0:
            return {'balance_score': 0, 'recommendation': 'No questions available'}
        
        actual_distribution = {level: count/total_questions for level, count in difficulty_distribution.items()}
        
        # Calculate balance score based on deviation from ideal
        total_deviation = 0
        for level in ideal_distribution:
            actual = actual_distribution.get(level, 0)
            ideal = ideal_distribution[level]
            deviation = abs(actual - ideal)
            total_deviation += deviation
        
        balance_score = max(0, 1 - total_deviation)
        
        # Generate recommendation
        if balance_score < 0.7:
            recommendation = "Consider adjusting difficulty distribution for better balance"
        else:
            recommendation = "Difficulty distribution is well-balanced"
        
        return {
            'balance_score': balance_score,
            'recommendation': recommendation,
            'actual_distribution': actual_distribution,
            'ideal_distribution': ideal_distribution
        }
    
    def _calculate_similarity_score(self, current_analysis: Dict, historical_data: Dict) -> float:
        """Calculate similarity score between current and historical data"""
        # This is a simplified similarity calculation
        # In practice, you might want to use more sophisticated methods
        
        score = 0.0
        factors = 0
        
        # Compare topic distributions
        if 'topic_weightage' in historical_data and 'topic_weightage' in current_analysis:
            current_topics = current_analysis.get('topic_weightage', {})
            historical_topics = historical_data.get('topic_weightage', {})
            
            common_topics = set(current_topics.keys()) & set(historical_topics.keys())
            if common_topics:
                topic_similarity = 0
                for topic in common_topics:
                    current_weight = current_topics[topic]
                    historical_weight = historical_topics[topic]
                    similarity = 1 - abs(current_weight - historical_weight) / max(current_weight, historical_weight, 1)
                    topic_similarity += similarity
                
                score += topic_similarity / len(common_topics)
                factors += 1
        
        # Compare type distributions
        if 'type_distribution' in historical_data and 'question_type_distribution' in current_analysis:
            current_types = current_analysis.get('question_type_distribution', {})
            historical_types = historical_data.get('type_distribution', {})
            
            common_types = set(current_types.keys()) & set(historical_types.keys())
            if common_types:
                type_similarity = 0
                for qtype in common_types:
                    current_count = current_types[qtype]
                    historical_count = historical_types[qtype]
                    similarity = 1 - abs(current_count - historical_count) / max(current_count, historical_count, 1)
                    type_similarity += similarity
                
                score += type_similarity / len(common_types)
                factors += 1
        
        return score / factors if factors > 0 else 0.0 