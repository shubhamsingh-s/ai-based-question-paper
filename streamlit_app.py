# Main application file for QuestVibe - Version 2.0
"import streamlit as st" 

import streamlit as st
import pandas as pd
import plotly.express as px
import random
import json
from datetime import datetime
import re
import sqlite3
import os
import requests
import openai
from typing import List, Dict, Any
# from advanced_analytics import advanced_analytics_dashboard
from collaboration_system import collaboration_dashboard
from streamlit_option_menu import option_menu
# from database_manager import (
#     create_connection,
#     create_tables,
# )
from transformers import pipeline
import torch
import streamlit_authenticator as stauth

# Page configuration
st.set_page_config(
    page_title="QuestVibe",
    page_icon="📚",
    layout="wide"
)

# Database setup
def init_database():
    conn = sqlite3.connect('user_data.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            institution TEXT NOT NULL,
            role TEXT DEFAULT 'user',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            session_start TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            session_end TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS question_generations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            subject TEXT,
            num_questions INTEGER,
            question_types TEXT,
            generated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    conn.commit()
    conn.close()

# Initialize database
init_database()

# Custom CSS for enhanced styling
st.markdown("""
<style>
/* Main background styling with aesthetic image */
.main {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    min-height: 100vh;
}
.stApp {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
    background-attachment: fixed;
}
.main .block-container {
    background: rgba(255, 255, 255, 0.2);
    backdrop-filter: blur(20px);
    border-radius: 20px;
    padding: 2rem;
    margin: 1rem;
    border: 1px solid rgba(255, 255, 255, 0.4);
    box-shadow: 0 8px 32px rgba(0,0,0,0.2);
}
h1 {
    background: linear-gradient(45deg, #FF6B6B, #4ECDC4, #45B7D1, #96CEB4);
    background-size: 300% 300%;
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    animation: gradient 3s ease infinite;
    text-align: center;
    font-size: 3.5rem;
    font-weight: 800;
    margin-bottom: 2rem;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
}
@keyframes gradient {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}
.stButton > button {
    background: linear-gradient(45deg, #667eea, #764ba2);
    color: white;
    border: none;
    border-radius: 15px;
    padding: 0.75rem 2rem;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: 0 4px 15px rgba(0,0,0,0.3);
}
.stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 20px rgba(0,0,0,0.4);
}
.stButton > button[data-baseweb="button"] {
    background: linear-gradient(45deg, #FF6B6B, #4ECDC4);
    font-size: 1.2rem;
    padding: 1rem 2.5rem;
    border-radius: 25px;
}
div[data-testid="column"] {
    background: rgba(255, 255, 255, 0.25);
    border-radius: 20px;
    padding: 2rem;
    margin: 1rem 0.5rem;
    box-shadow: 0 8px 32px rgba(0,0,0,0.15);
    border: 1px solid rgba(255, 255, 255, 0.4);
    transition: all 0.3s ease;
    backdrop-filter: blur(20px);
}
div[data-testid="column"]:hover {
    transform: translateY(-5px);
    box-shadow: 0 12px 40px rgba(0,0,0,0.2);
    background: rgba(255, 255, 255, 0.3);
}
div[data-testid="metric-container"] {
    background: linear-gradient(135deg, rgba(102, 126, 234, 0.9), rgba(118, 75, 162, 0.9));
    border-radius: 15px;
    padding: 1.5rem;
    color: white;
    text-align: center;
    box-shadow: 0 4px 15px rgba(0,0,0,0.3);
    margin: 0.5rem 0;
    backdrop-filter: blur(10px);
}
.footer {
    background: linear-gradient(135deg, rgba(44, 62, 80, 0.9), rgba(52, 73, 94, 0.9));
    color: white;
    text-align: center;
    padding: 2rem;
    border-radius: 20px;
    margin-top: 3rem;
    box-shadow: 0 8px 32px rgba(0,0,0,0.2);
    backdrop-filter: blur(15px);
}
p, h2, h3, h4, h5, h6 {
    color: #ffffff;
    font-weight: 600;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.5);
}
div[data-testid="stMarkdown"] {
    background: rgba(255, 255, 255, 0.15);
    padding: 1rem;
    border-radius: 10px;
    margin: 0.5rem 0;
    backdrop-filter: blur(15px);
    border: 1px solid rgba(255, 255, 255, 0.3);
}
.stTextInput > div > div > input {
    background: rgba(255, 255, 255, 0.25);
    border-radius: 10px;
    border: 1px solid rgba(255, 255, 255, 0.4);
    color: #000000 !important;
    backdrop-filter: blur(15px);
    font-weight: 500;
}
.stTextInput > div > div > input::placeholder {
    color: rgba(0, 0, 0, 0.6);
}
/* Make sure text is visible in all input fields */
input, textarea, select {
    color: #000000 !important;
}
/* Ensure password field text is also visible */
input[type="password"] {
    color: #000000 !important;
}
.stSelectbox > div > div > div {
    background: rgba(255, 255, 255, 0.25);
    border-radius: 10px;
    border: 1px solid rgba(255, 255, 255, 0.4);
    backdrop-filter: blur(15px);
}
.stSlider > div > div > div > div {
    background: rgba(255, 255, 255, 0.25);
}
.stAlert {
    background: rgba(255, 255, 255, 0.2);
    border-radius: 15px;
    backdrop-filter: blur(15px);
    border: 1px solid rgba(255, 255, 255, 0.4);
}
.streamlit-expanderHeader {
    background: rgba(255, 255, 255, 0.2);
    border-radius: 10px;
    padding: 1rem;
    margin: 0.5rem 0;
    border: 1px solid rgba(255, 255, 255, 0.4);
    backdrop-filter: blur(15px);
    color: #ffffff;
    font-weight: 600;
}
.streamlit-expanderContent {
    background: rgba(255, 255, 255, 0.15);
    border-radius: 10px;
    padding: 1rem;
    margin: 0.5rem 0;
    border: 1px solid rgba(255, 255, 255, 0.3);
    backdrop-filter: blur(15px);
}
.stTabs [data-baseweb="tab-list"] {
    background: rgba(255, 255, 255, 0.2);
    border-radius: 10px;
    padding: 0.5rem;
    margin: 0.5rem 0;
}
</style>
""", unsafe_allow_html=True)

# --- Simple user database (for demo) ---
USERS = {
    "admin": {"password": "admin123", "role": "admin", "name": "Administrator"},
    "teacher": {"password": "teacher123", "role": "teacher", "name": "Teacher"},
    "student": {"password": "student123", "role": "student", "name": "Student"},
    "demo": {"password": "demo123", "role": "teacher", "name": "Demo User"}
}

# Super User credentials (hidden admin access)
SUPER_USER = {
    "username": "superadmin",
    "password": "questvibe2024",
    "role": "super_admin",
    "name": "Super Administrator"
}

def save_user_to_database(name, institution):
    conn = sqlite3.connect('user_data.db')
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO users (name, institution, role) 
        VALUES (?, ?, ?)
    ''', (name, institution, 'user'))
    user_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return user_id

def get_user_by_id(user_id):
    """Get user by ID"""
    conn = sqlite3.connect('user_data.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))
    user = cursor.fetchone()
    conn.close()
    if user:
        return dict(user)
    return None

def log_session(user_id):
    conn = sqlite3.connect('user_data.db')
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO sessions (user_id) 
        VALUES (?)
    ''', (user_id,))
    session_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return session_id

def log_question_generation(user_id, subject, num_questions, question_types):
    conn = sqlite3.connect('user_data.db')
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO question_generations (user_id, subject, num_questions, question_types) 
        VALUES (?, ?, ?, ?)
    ''', (user_id, subject, num_questions, json.dumps(question_types)))
    conn.commit()
    conn.close()

def get_database_stats():
    conn = sqlite3.connect('user_data.db')
    cursor = conn.cursor()
    
    # Get user count
    cursor.execute('SELECT COUNT(*) FROM users')
    user_count = cursor.fetchone()[0]
    
    # Get total generations
    cursor.execute('SELECT COUNT(*) FROM question_generations')
    generation_count = cursor.fetchone()[0]
    
    # Get recent users
    cursor.execute('SELECT name, institution, created_at FROM users ORDER BY created_at DESC LIMIT 10')
    recent_users = cursor.fetchall()
    
    # Get recent generations
    cursor.execute('''
        SELECT u.name, u.institution, qg.subject, qg.num_questions, qg.generated_at 
        FROM question_generations qg 
        JOIN users u ON qg.user_id = u.id 
        ORDER BY qg.generated_at DESC LIMIT 10
    ''')
    recent_generations = cursor.fetchall()
    
    conn.close()
    return user_count, generation_count, recent_users, recent_generations

def admin_dashboard():
    st.markdown("## 📊 Admin Dashboard")
    st.write("Database analytics and user information")
    
    user_count, generation_count, recent_users, recent_generations = get_database_stats()
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Users", user_count, "👥")
    
    with col2:
        st.metric("Total Generations", generation_count, "📝")
    
    with col3:
        st.metric("Active Sessions", "Live", "🟢")
    
    with col4:
        st.metric("Database Status", "Connected", "✅")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 👥 Recent Users")
        if recent_users:
            for user in recent_users:
                st.write(f"**{user[0]}** - {user[1]} ({user[2][:10]}...)")
        else:
            st.write("No users yet")
    
    with col2:
        st.markdown("### 📝 Recent Generations")
        if recent_generations:
            for gen in recent_generations:
                st.write(f"**{gen[0]}** - {gen[2]} ({gen[3]} questions)")
        else:
            st.write("No generations yet")
    
    if st.button("🔙 Back to Dashboard"):
        st.session_state.show_admin = False
        st.rerun()

def super_admin_dashboard():
    st.markdown("## 🔓 Super Admin Dashboard")
    st.warning("⚠️ **SUPER ADMIN ACCESS** - Full database control")
    
    # Get all database data
    conn = sqlite3.connect('user_data.db')
    
    # Users table
    st.markdown("### 👥 All Users")
    users_df = pd.read_sql_query("SELECT * FROM users ORDER BY created_at DESC", conn)
    if not users_df.empty:
        st.dataframe(users_df, use_container_width=True)
        
        # Download users data
        csv = users_df.to_csv(index=False)
        st.download_button(
            label="📥 Download Users Data (CSV)",
            data=csv,
            file_name="users_data.csv",
            mime="text/csv"
        )
    else:
        st.write("No users found in database")
    
    # Sessions table
    st.markdown("### 📊 User Sessions")
    sessions_df = pd.read_sql_query("""
        SELECT s.*, u.name, u.institution 
        FROM sessions s 
        JOIN users u ON s.user_id = u.id 
        ORDER BY s.session_start DESC
    """, conn)
    if not sessions_df.empty:
        st.dataframe(sessions_df, use_container_width=True)
        
        csv = sessions_df.to_csv(index=False)
        st.download_button(
            label="📥 Download Sessions Data (CSV)",
            data=csv,
            file_name="sessions_data.csv",
            mime="text/csv"
        )
    else:
        st.write("No sessions found in database")
    
    # Question generations table
    st.markdown("### 📝 Question Generations")
    generations_df = pd.read_sql_query("""
        SELECT qg.*, u.name, u.institution 
        FROM question_generations qg 
        JOIN users u ON qg.user_id = u.id 
        ORDER BY qg.generated_at DESC
    """, conn)
    if not generations_df.empty:
        st.dataframe(generations_df, use_container_width=True)
        
        csv = generations_df.to_csv(index=False)
        st.download_button(
            label="📥 Download Generations Data (CSV)",
            data=csv,
            file_name="generations_data.csv",
            mime="text/csv"
        )
    else:
        st.write("No question generations found in database")
    
    # Database statistics
    st.markdown("### 📈 Database Statistics")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_users = len(users_df)
        st.metric("Total Users", total_users, "👥")
    
    with col2:
        total_sessions = len(sessions_df)
        st.metric("Total Sessions", total_sessions, "📊")
    
    with col3:
        total_generations = len(generations_df)
        st.metric("Total Generations", total_generations, "📝")
    
    with col4:
        if not generations_df.empty:
            avg_questions = generations_df['num_questions'].mean()
            st.metric("Avg Questions/Gen", f"{avg_questions:.1f}", "❓")
        else:
            st.metric("Avg Questions/Gen", "0", "❓")
    
    # Charts
    if not generations_df.empty:
        st.markdown("### 📊 Analytics Charts")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Subject distribution
            subject_counts = generations_df['subject'].value_counts()
            fig1 = px.pie(values=subject_counts.values, names=subject_counts.index, title="Question Generations by Subject")
            st.plotly_chart(fig1, use_container_width=True)
        
        with col2:
            # Questions per generation
            fig2 = px.histogram(generations_df, x='num_questions', title="Distribution of Questions per Generation")
            st.plotly_chart(fig2, use_container_width=True)
    
    # Database management
    st.markdown("### 🗄️ Database Management")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🗑️ Clear All Data", type="secondary"):
            if st.checkbox("I understand this will delete ALL data permanently"):
                cursor = conn.cursor()
                cursor.execute("DELETE FROM question_generations")
                cursor.execute("DELETE FROM sessions")
                cursor.execute("DELETE FROM users")
                conn.commit()
                st.success("🗑️ All data cleared!")
                st.rerun()
    
    with col2:
        if st.button("📊 Export Full Database", type="secondary"):
            # Create a zip file with all data
            import zipfile
            import io
            
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                # Add each table as CSV
                for table_name, df in [("users", users_df), ("sessions", sessions_df), ("generations", generations_df)]:
                    csv_data = df.to_csv(index=False)
                    zip_file.writestr(f"{table_name}.csv", csv_data)
            
            st.download_button(
                label="📦 Download Full Database (ZIP)",
                data=zip_buffer.getvalue(),
                file_name="questvibe_database.zip",
                mime="application/zip"
            )
    
    # AI Database Management
    st.markdown("### 🤖 AI Database Management")
    
    # Get database status
    db_status = st.session_state.questvibe_ai_db.get_database_status()
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Subjects", db_status['total_subjects'], "📚")
    
    with col2:
        st.metric("Total Topics", db_status['total_topics'], "🎯")
    
    with col3:
        st.metric("Duplicates Found", len(db_status['duplicates']), "⚠️")
    
    with col4:
        if st.button("🧹 Clean Duplicates", type="secondary"):
            st.session_state.questvibe_ai_db.cleanup_database()
            st.success("✅ Database cleaned!")
            st.rerun()
    
    # Show branch statistics
    st.markdown("**📊 Subjects by Branch:**")
    for branch, count in db_status['branch_counts']:
        st.write(f"• **{branch}**: {count} subjects")
    
    # Show duplicates if any
    if db_status['duplicates']:
        st.markdown("**⚠️ Duplicate Subjects Found:**")
        for branch, subject, count in db_status['duplicates']:
            st.write(f"• {branch} - {subject} (appears {count} times)")
    
    # Repopulate database option
    st.markdown("**🔄 Database Operations:**")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🔄 Repopulate AI Database", type="primary"):
            st.session_state.questvibe_ai_db.populate_database()
            st.success("✅ AI Database repopulated!")
            st.rerun()
    
    with col2:
        if st.button("📊 Refresh Status", type="secondary"):
            st.rerun()
    
    # ChatGPT Configuration
    st.markdown("### 🤖 ChatGPT Configuration")
    
    # API Key management
    with st.expander("🔑 ChatGPT API Configuration", expanded=False):
        api_key = st.text_input("OpenAI API Key", type="password", 
                               help="Enter your OpenAI API key to enable ChatGPT integration")
        
        if api_key:
            st.session_state.questvibe_chatgpt.api_key = api_key
            st.success("✅ API key configured!")
        
        # Test ChatGPT connection
        if st.button("🧪 Test ChatGPT Connection", type="secondary"):
            if st.session_state.questvibe_chatgpt.api_key:
                try:
                    test_questions = st.session_state.questvibe_chatgpt.generate_questions(
                        "Test Subject", ["Test Topic"], 1, ["MCQ"]
                    )
                    if test_questions:
                        st.success("✅ ChatGPT connection successful!")
                        st.write("**Sample Question:**")
                        st.write(test_questions[0]['question'])
                    else:
                        st.error("❌ ChatGPT test failed - no questions generated")
                except Exception as e:
                    st.error(f"❌ ChatGPT test failed: {str(e)}")
            else:
                st.warning("⚠️ Please enter an API key first")
    
    # ChatGPT Statistics
    st.markdown("**📊 ChatGPT Usage Statistics:**")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        api_status = "✅ Connected" if st.session_state.questvibe_chatgpt.api_key else "❌ Not Connected"
        st.metric("API Status", api_status, "🔗")
    
    with col2:
        st.metric("Model", st.session_state.questvibe_chatgpt.model, "🤖")
    
    with col3:
        st.metric("Max Tokens", st.session_state.questvibe_chatgpt.max_tokens, "📝")
    
    # ChatGPT Features
    st.markdown("**✨ ChatGPT Features:**")
    features = [
        "Intelligent question generation",
        "Multiple difficulty levels",
        "Bloom's taxonomy integration", 
        "Industry-relevant content",
        "JSON response parsing",
        "Fallback generation system"
    ]
    
    for feature in features:
        st.write(f"• {feature}")
    
    conn.close()
    
    if st.button("🔙 Back to Dashboard"):
        st.session_state.show_super_admin = False
        st.rerun()

# AI Database Population System
class QuestVibeAIDatabase:
    def __init__(self):
        self.engineering_subjects = {
            "Computer Science Engineering": {
                "subjects": [
                    "Data Structures and Algorithms",
                    "Object-Oriented Programming",
                    "Database Management Systems",
                    "Computer Networks",
                    "Operating Systems",
                    "Software Engineering",
                    "Web Technologies",
                    "Machine Learning",
                    "Artificial Intelligence",
                    "Computer Architecture",
                    "Digital Logic Design",
                    "Computer Organization",
                    "Data Communication",
                    "Information Security",
                    "Cloud Computing",
                    "Big Data Analytics",
                    "Internet of Things",
                    "Mobile Computing",
                    "Computer Graphics",
                    "System Programming"
                ],
                "exam_topics": {
                    "Data Structures and Algorithms": [
                        "Arrays and Linked Lists", "Stacks and Queues", "Trees and Graphs",
                        "Sorting Algorithms", "Searching Algorithms", "Dynamic Programming",
                        "Greedy Algorithms", "Complexity Analysis", "Recursion", "Hash Tables"
                    ],
                    "Database Management Systems": [
                        "ER Model", "Relational Model", "SQL Queries", "Normalization",
                        "Transaction Management", "Concurrency Control", "Indexing",
                        "Database Security", "Distributed Databases", "NoSQL Databases"
                    ],
                    "Computer Networks": [
                        "OSI Model", "TCP/IP Protocol", "Network Topologies", "Routing Algorithms",
                        "Network Security", "Wireless Networks", "Network Management",
                        "Quality of Service", "Network Protocols", "Internet Architecture"
                    ]
                }
            },
            "Information Technology": {
                "subjects": [
                    "Programming Fundamentals",
                    "Data Structures",
                    "Database Systems",
                    "Computer Networks",
                    "Web Development",
                    "Software Testing",
                    "Information Security",
                    "Cloud Computing",
                    "Mobile App Development",
                    "Data Analytics",
                    "System Analysis and Design",
                    "Operating Systems",
                    "Computer Architecture",
                    "Digital Electronics",
                    "Computer Graphics",
                    "Multimedia Systems",
                    "E-Commerce",
                    "Enterprise Systems",
                    "Network Administration",
                    "IT Project Management"
                ],
                "exam_topics": {
                    "Programming Fundamentals": [
                        "Variables and Data Types", "Control Structures", "Functions",
                        "Arrays and Strings", "Object-Oriented Concepts", "Exception Handling",
                        "File I/O", "Memory Management", "Debugging", "Code Optimization"
                    ],
                    "Web Development": [
                        "HTML and CSS", "JavaScript", "PHP", "ASP.NET", "Web Servers",
                        "Client-Side Scripting", "Server-Side Scripting", "Web Security",
                        "Responsive Design", "Web APIs"
                    ]
                }
            },
            "Electronics and Communication": {
                "subjects": [
                    "Digital Electronics",
                    "Analog Electronics",
                    "Communication Systems",
                    "Signals and Systems",
                    "Electromagnetic Theory",
                    "Antenna Theory",
                    "Microwave Engineering",
                    "Optical Communication",
                    "Satellite Communication",
                    "Wireless Communication",
                    "Digital Signal Processing",
                    "VLSI Design",
                    "Microprocessors",
                    "Control Systems",
                    "Power Electronics",
                    "Electronic Devices",
                    "Circuit Theory",
                    "Network Analysis",
                    "Telecommunication Networks",
                    "Information Theory"
                ],
                "exam_topics": {
                    "Digital Electronics": [
                        "Boolean Algebra", "Logic Gates", "Combinational Circuits",
                        "Sequential Circuits", "Flip-Flops", "Counters", "Registers",
                        "Memory Devices", "Digital ICs", "VHDL Programming"
                    ],
                    "Communication Systems": [
                        "Amplitude Modulation", "Frequency Modulation", "Digital Modulation",
                        "Pulse Modulation", "Multiplexing", "Demodulation", "Noise Analysis",
                        "Channel Coding", "Error Detection", "Information Theory"
                    ]
                }
            },
            "Mechanical Engineering": {
                "subjects": [
                    "Engineering Mechanics",
                    "Strength of Materials",
                    "Machine Design",
                    "Thermodynamics",
                    "Fluid Mechanics",
                    "Heat Transfer",
                    "Manufacturing Processes",
                    "CAD/CAM",
                    "Automobile Engineering",
                    "Robotics",
                    "Control Systems",
                    "Material Science",
                    "Dynamics of Machines",
                    "Theory of Machines",
                    "Industrial Engineering",
                    "Refrigeration and Air Conditioning",
                    "Power Plant Engineering",
                    "Automation",
                    "Quality Control",
                    "Project Management"
                ],
                "exam_topics": {
                    "Engineering Mechanics": [
                        "Statics", "Dynamics", "Kinematics", "Kinetics", "Friction",
                        "Centroids", "Moment of Inertia", "Work and Energy", "Impulse and Momentum",
                        "Vibration Analysis"
                    ],
                    "Thermodynamics": [
                        "First Law", "Second Law", "Entropy", "Gas Cycles", "Vapor Cycles",
                        "Heat Engines", "Refrigeration Cycles", "Psychrometrics", "Combustion",
                        "Steam Tables"
                    ]
                }
            },
            "Civil Engineering": {
                "subjects": [
                    "Structural Analysis",
                    "Reinforced Concrete Design",
                    "Steel Structures",
                    "Transportation Engineering",
                    "Geotechnical Engineering",
                    "Hydraulics",
                    "Environmental Engineering",
                    "Surveying",
                    "Construction Management",
                    "Highway Engineering",
                    "Bridge Engineering",
                    "Water Resources",
                    "Irrigation Engineering",
                    "Town Planning",
                    "Building Materials",
                    "Soil Mechanics",
                    "Foundation Engineering",
                    "Earthquake Engineering",
                    "Traffic Engineering",
                    "Project Planning"
                ],
                "exam_topics": {
                    "Structural Analysis": [
                        "Beam Analysis", "Truss Analysis", "Frame Analysis", "Influence Lines",
                        "Deflection Analysis", "Stability Analysis", "Matrix Methods",
                        "Finite Element Method", "Plastic Analysis", "Dynamic Analysis"
                    ],
                    "Reinforced Concrete Design": [
                        "Flexural Design", "Shear Design", "Column Design", "Footing Design",
                        "Slab Design", "Beam Design", "Detailing", "Durability", "Serviceability",
                        "Code Provisions"
                    ]
                }
            }
        }
    
    def populate_database(self):
        """Automatically populate database with engineering subjects and topics"""
        conn = sqlite3.connect('user_data.db')
        cursor = conn.cursor()
        
        # Create subjects table if not exists with UNIQUE constraint
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS subjects (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                branch TEXT NOT NULL,
                subject_name TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(branch, subject_name)
            )
        ''')
        
        # Create topics table if not exists
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS topics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                subject_id INTEGER,
                topic_name TEXT NOT NULL,
                difficulty_level TEXT DEFAULT 'Medium',
                bloom_level TEXT DEFAULT 'Understanding',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(subject_id, topic_name),
                FOREIGN KEY (subject_id) REFERENCES subjects (id)
            )
        ''')
        
        # Create exam_patterns table if not exists
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS exam_patterns (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                subject_id INTEGER,
                pattern_name TEXT NOT NULL,
                total_questions INTEGER,
                mcq_count INTEGER,
                short_answer_count INTEGER,
                long_answer_count INTEGER,
                case_study_count INTEGER,
                time_duration INTEGER,
                marks_per_question REAL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (subject_id) REFERENCES subjects (id)
            )
        ''')
        
        # Clear existing data to prevent duplicates
        cursor.execute('DELETE FROM topics')
        cursor.execute('DELETE FROM subjects')
        
        # Populate subjects and topics
        for branch, data in self.engineering_subjects.items():
            for subject in data["subjects"]:
                try:
                    # Insert subject with UNIQUE constraint
                    cursor.execute('''
                        INSERT INTO subjects (branch, subject_name) 
                        VALUES (?, ?)
                    ''', (branch, subject))
                    
                    subject_id = cursor.lastrowid
                    
                    # Insert topics if available
                    if subject in data.get("exam_topics", {}):
                        for topic in data["exam_topics"][subject]:
                            try:
                                cursor.execute('''
                                    INSERT INTO topics (subject_id, topic_name) 
                                    VALUES (?, ?)
                                ''', (subject_id, topic))
                            except sqlite3.IntegrityError:
                                # Topic already exists, skip
                                pass
                except sqlite3.IntegrityError:
                    # Subject already exists, get its ID and add topics
                    cursor.execute('SELECT id FROM subjects WHERE branch = ? AND subject_name = ?', (branch, subject))
                    result = cursor.fetchone()
                    if result:
                        subject_id = result[0]
                        # Insert topics if available
                        if subject in data.get("exam_topics", {}):
                            for topic in data["exam_topics"][subject]:
                                try:
                                    cursor.execute('''
                                        INSERT INTO topics (subject_id, topic_name) 
                                        VALUES (?, ?)
                                    ''', (subject_id, topic))
                                except sqlite3.IntegrityError:
                                    # Topic already exists, skip
                                    pass
        
        # Clear and insert default exam patterns
        cursor.execute('DELETE FROM exam_patterns')
        default_patterns = [
            ("Mid Semester", 20, 10, 5, 3, 2, 90, 5.0),
            ("End Semester", 30, 15, 8, 5, 2, 180, 5.0),
            ("Unit Test", 15, 8, 4, 2, 1, 60, 5.0)
        ]
        
        for pattern_name, total, mcq, short, long, case_study, duration, marks in default_patterns:
            cursor.execute('''
                INSERT INTO exam_patterns 
                (pattern_name, total_questions, mcq_count, short_answer_count, long_answer_count, case_study_count, time_duration, marks_per_question) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (pattern_name, total, mcq, short, long, case_study, duration, marks))
        
        conn.commit()
        conn.close()
        
        return True
    
    def get_subjects_by_branch(self, branch):
        """Get all subjects for a specific engineering branch"""
        conn = sqlite3.connect('user_data.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT subject_name FROM subjects 
            WHERE branch = ? 
            ORDER BY subject_name
        ''', (branch,))
        
        subjects = [row[0] for row in cursor.fetchall()]
        conn.close()
        return subjects
    
    def get_topics_by_subject(self, subject_name):
        """Get all topics for a specific subject"""
        conn = sqlite3.connect('user_data.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT t.topic_name FROM topics t
            JOIN subjects s ON t.subject_id = s.id
            WHERE s.subject_name = ?
            ORDER BY t.topic_name
        ''', (subject_name,))
        
        topics = [row[0] for row in cursor.fetchall()]
        conn.close()
        return topics
    
    def get_exam_patterns(self):
        """Get all available exam patterns"""
        conn = sqlite3.connect('user_data.db')
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM exam_patterns ORDER BY pattern_name')
        patterns = cursor.fetchall()
        conn.close()
        return patterns

    def cleanup_database(self):
        """Clean up duplicate data and ensure database integrity"""
        conn = sqlite3.connect('user_data.db')
        cursor = conn.cursor()
        
        # Remove duplicate subjects (keep the first occurrence)
        cursor.execute('''
            DELETE FROM subjects 
            WHERE id NOT IN (
                SELECT MIN(id) 
                FROM subjects 
                GROUP BY branch, subject_name
            )
        ''')
        
        # Remove duplicate topics (keep the first occurrence)
        cursor.execute('''
            DELETE FROM topics 
            WHERE id NOT IN (
                SELECT MIN(id) 
                FROM topics 
                GROUP BY subject_id, topic_name
            )
        ''')
        
        # Remove orphaned topics (topics without subjects)
        cursor.execute('''
            DELETE FROM topics 
            WHERE subject_id NOT IN (SELECT id FROM subjects)
        ''')
        
        conn.commit()
        conn.close()
        
        return True
    
    def get_database_status(self):
        """Get database status and statistics"""
        conn = sqlite3.connect('user_data.db')
        cursor = conn.cursor()
        
        # Count subjects by branch
        cursor.execute('''
            SELECT branch, COUNT(*) as count 
            FROM subjects 
            GROUP BY branch 
            ORDER BY branch
        ''')
        branch_counts = cursor.fetchall()
        
        # Count total topics
        cursor.execute('SELECT COUNT(*) FROM topics')
        total_topics = cursor.fetchone()[0]
        
        # Count total subjects
        cursor.execute('SELECT COUNT(*) FROM subjects')
        total_subjects = cursor.fetchone()[0]
        
        # Check for duplicates
        cursor.execute('''
            SELECT branch, subject_name, COUNT(*) as count 
            FROM subjects 
            GROUP BY branch, subject_name 
            HAVING COUNT(*) > 1
        ''')
        duplicates = cursor.fetchall()
        
        conn.close()
        
        return {
            'branch_counts': branch_counts,
            'total_topics': total_topics,
            'total_subjects': total_subjects,
            'duplicates': duplicates
        }

# Initialize AI Database System
if 'questvibe_ai_db' not in st.session_state:
    st.session_state.questvibe_ai_db = QuestVibeAIDatabase()
    # Automatically populate database on first run
    st.session_state.questvibe_ai_db.populate_database()

# ChatGPT Integration
class QuestVibeChatGPT:
    def __init__(self, api_key=None):
        self.api_key = api_key or st.secrets.get("OPENAI_API_KEY", "")
        self.model = "gpt-3.5-turbo"
        self.max_tokens = 1000
        
    def generate_questions(self, subject: str, topics: List[str], num_questions: int, question_types: List[str]) -> List[Dict]:
        """Generate questions using ChatGPT with fallback"""
        prompt = self.create_question_prompt(subject, topics, num_questions, question_types)
        response = self.call_chatgpt_api(prompt)
        
        if response:
            return self.parse_chatgpt_response(response, question_types)
        else:
            return self.generate_fallback_questions(subject, topics, num_questions, question_types)
    
    def create_question_prompt(self, subject: str, topics: List[str], num_questions: int, question_types: List[str]) -> str:
        """Create a detailed prompt for ChatGPT"""
        topics_text = ", ".join(topics)
        types_text = ", ".join(question_types)
        
        prompt = f"""
        You are an expert educator creating questions for {subject}. 
        
        Available topics: {topics_text}
        Question types needed: {types_text}
        Number of questions: {num_questions}
        
        Please generate {num_questions} questions with the following requirements:
        
        1. Questions should be relevant to the topics provided
        2. Mix of question types: {types_text}
        3. For MCQs, provide 4 options (A, B, C, D) with one correct answer
        4. Questions should test different cognitive levels (Remember, Understand, Apply, Analyze, Evaluate, Create)
        5. Difficulty should vary from easy to hard
        6. Questions should be practical and industry-relevant
        
        Format your response as JSON:
        {{
            "questions": [
                {{
                    "type": "MCQ/Short Answer/Long Answer/Case Study",
                    "question": "Question text here",
                    "options": ["A", "B", "C", "D"] (only for MCQ),
                    "correct_answer": "A" (only for MCQ),
                    "difficulty": "Easy/Medium/Hard",
                    "bloom_level": "Remember/Understand/Apply/Analyze/Evaluate/Create",
                    "topic": "specific topic from the list"
                }}
            ]
        }}
        
        Make sure the questions are high-quality, educational, and suitable for engineering students.
        """
        
        return prompt
    
    def call_chatgpt_api(self, prompt: str) -> str:
        """Call ChatGPT API using requests"""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "gpt-3.5-turbo",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
        
        try:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=data
            )
            
            if response.status_code == 200:
                return response.json()["choices"][0]["message"]["content"]
            else:
                raise Exception(f"API call failed: {response.status_code}")
        except Exception as e:
            st.error(f"An error occurred while calling ChatGPT API: {e}")
            return ""
    
    def parse_chatgpt_response(self, response: str, question_types: List[str]) -> List[Dict]:
        """Parse the response from ChatGPT API"""
        try:
            # Try to extract JSON from response
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            
            if json_start != -1 and json_end != 0:
                json_str = response[json_start:json_end]
                data = json.loads(json_str)
                
                if "questions" in data:
                    return data["questions"]
            
            # Fallback: parse manually
            return self.parse_manual_response(response, question_types)
            
        except Exception as e:
            # Only show parsing error to super admins
            if hasattr(st.session_state, 'current_user') and st.session_state.current_user and st.session_state.current_user.get('role') == 'super_admin':
                st.warning(f"Failed to parse ChatGPT response: {str(e)}")
            return self.parse_manual_response(response, question_types)
    
    def parse_manual_response(self, response: str, question_types: List[str]) -> List[Dict]:
        """Manually parse response if JSON parsing fails"""
        questions = []
        lines = response.split('\n')
        current_question = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Look for question patterns
            if any(qtype.lower() in line.lower() for qtype in question_types):
                if current_question:
                    questions.append(current_question)
                
                current_question = {
                    "type": "MCQ" if "mcq" in line.lower() else "Short Answer",
                    "question": line,
                    "options": [],
                    "difficulty": "Medium",
                    "bloom_level": "Understand",
                    "topic": "General"
                }
            elif current_question and line.startswith(('A.', 'B.', 'C.', 'D.')):
                current_question["options"].append(line)
        
        if current_question:
            questions.append(current_question)
        
        return questions
    
    def generate_fallback_questions(self, subject: str, topics: List[str], num_questions: int, question_types: List[str]) -> List[Dict]:
        """Generate questions without ChatGPT API"""
        questions = []
        
        for i in range(num_questions):
            topic = random.choice(topics)
            q_type = random.choice(question_types)
            
            if q_type == "MCQ":
                question = {
                    "type": "MCQ",
                    "question": f"{i+1}. What is {topic}?",
                    "options": [
                        f"A. {topic} is a fundamental concept",
                        f"B. {topic} is an advanced technique", 
                        f"C. {topic} is a basic principle",
                        f"D. {topic} is a complex system"
                    ],
                    "correct_answer": "A",
                    "difficulty": random.choice(["Easy", "Medium", "Hard"]),
                    "bloom_level": "Understand",
                    "topic": topic
                }
            else:
                question = {
                    "type": q_type,
                    "question": f"{i+1}. Explain {topic} in detail and discuss its applications in {subject}.",
                    "difficulty": random.choice(["Easy", "Medium", "Hard"]),
                    "bloom_level": "Analyze",
                    "topic": topic
                }
            
            questions.append(question)
        
        return questions
    
    def analyze_question_quality(self, questions: List[Dict]) -> Dict[str, Any]:
        """Analyze the quality and distribution of generated questions"""
        analysis = {
            "total_questions": len(questions),
            "type_distribution": {},
            "difficulty_distribution": {},
            "bloom_distribution": {},
            "topic_coverage": set(),
            "quality_score": 0
        }
        
        for q in questions:
            # Type distribution
            q_type = q.get("type", "Unknown")
            analysis["type_distribution"][q_type] = analysis["type_distribution"].get(q_type, 0) + 1
            
            # Difficulty distribution
            difficulty = q.get("difficulty", "Medium")
            analysis["difficulty_distribution"][difficulty] = analysis["difficulty_distribution"].get(difficulty, 0) + 1
            
            # Bloom level distribution
            bloom = q.get("bloom_level", "Understand")
            analysis["bloom_distribution"][bloom] = analysis["bloom_distribution"].get(bloom, 0) + 1
            
            # Topic coverage
            topic = q.get("topic", "General")
            analysis["topic_coverage"].add(topic)
        
        # Calculate quality score
        analysis["quality_score"] = min(100, len(questions) * 2 + len(analysis["topic_coverage"]) * 5)
        analysis["topic_coverage"] = list(analysis["topic_coverage"])
        
        return analysis

# Initialize ChatGPT
if 'questvibe_chatgpt' not in st.session_state:
    st.session_state.questvibe_chatgpt = QuestVibeChatGPT()

# --- KeyBERT-powered topic extraction ---
from keybert import KeyBERT
import spacy

# Load models globally for efficiency
if 'kw_model' not in st.session_state:
    st.session_state.kw_model = KeyBERT(model='all-MiniLM-L6-v2')
if 'spacy_nlp' not in st.session_state:
    st.session_state.spacy_nlp = spacy.load("en_core_web_sm")

def keybert_syllabus_parser(syllabus_text, max_keywords=20):
    kw_model = st.session_state.kw_model
    nlp = st.session_state.spacy_nlp
    # 1. KeyBERT: semantic keyword/keyphrase extraction
    keybert_keywords = kw_model.extract_keywords(
        syllabus_text,
        keyphrase_ngram_range=(1, 4),
        stop_words='english',
        top_n=max_keywords
    )
    keybert_phrases = set([kw for kw, score in keybert_keywords])
    # 2. Noun phrase extraction (spaCy)
    doc = nlp(syllabus_text)
    noun_phrases = set(chunk.text.strip() for chunk in doc.noun_chunks if len(chunk.text.split()) < 8)
    # 3. Combine and deduplicate
    all_topics = keybert_phrases | noun_phrases
    all_topics = [t for t in all_topics if len(t) > 2 and not t.isdigit()]
    return list(all_topics)

def main_dashboard():
    """The main dashboard shown after the user logs in."""
    st.sidebar.title("Menu")
    user_role = st.session_state.current_user.get('role', 'user')
    # Build sidebar options based on user role
    sidebar_options = [
        "Upload Previous Paper",
        "Upload Syllabus",
        "Auto Question Paper",
        "Predictive Analysis",
        "Collaboration"
    ]
    sidebar_icons = [
        "file-earmark-arrow-up",
        "file-earmark-text",
        "robot",
        "bi bi-microscope",
        "people-fill"
    ]
    if user_role in ["admin", "super_admin"]:
        sidebar_options.append("Admin")
        sidebar_icons.append("shield-lock-fill")
    if user_role == "super_admin":
        sidebar_options.append("Super Admin")
        sidebar_icons.append("key-fill")
    with st.sidebar:
        if 'active_dashboard' not in st.session_state:
            st.session_state.active_dashboard = sidebar_options[0]
        selected = option_menu(
            menu_title=None,
            options=sidebar_options,
            icons=sidebar_icons,
            menu_icon="cast",
            default_index=0
        )
        st.session_state.active_dashboard = selected
        if st.button("Logout"):
            st.session_state.current_user = None
            st.rerun()
    st.markdown(f"### Welcome, {st.session_state.current_user['name']}!")

    if st.session_state.active_dashboard == "Predictive Analysis":
        st.markdown("#### 🔬 Predictive Analysis (ML-powered)")
        st.info("Upload previous year papers and syllabus to predict the most probable questions and classify them by Bloom's Taxonomy.")
        st.markdown("**Step 1:** Upload previous year papers (PDF, DOCX, TXT or paste text)")
        prev_file = st.file_uploader("Upload Previous Year Paper", type=["pdf", "docx", "txt"], key="predict_prev_upload")
        prev_text = ""
        if prev_file is not None:
            if prev_file.type == "application/pdf":
                import PyPDF2
                reader = PyPDF2.PdfReader(prev_file)
                prev_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif prev_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                import docx
                doc = docx.Document(prev_file)
                prev_text = "\n".join([para.text for para in doc.paragraphs])
            elif prev_file.type == "text/plain":
                prev_text = prev_file.read().decode("utf-8")
        prev_text_input = st.text_area("Or paste previous paper text here", height=150, key="predict_prev_text")
        prev_content = prev_text or prev_text_input
        st.session_state.predict_prev_content = prev_content
        if prev_content:
            st.markdown("**Preview: Previous Paper Content**")
            st.write(prev_content[:1000] + ("..." if len(prev_content) > 1000 else ""))
        else:
            st.info("Upload a file or paste text to see the content here.")

        st.markdown("**Step 2:** Upload syllabus (PDF, DOCX, TXT or paste text)")
        syllabus_file = st.file_uploader("Upload Syllabus", type=["pdf", "docx", "txt"], key="predict_syllabus_upload")
        syllabus_text = ""
        if syllabus_file is not None:
            if syllabus_file.type == "application/pdf":
                import PyPDF2
                reader = PyPDF2.PdfReader(syllabus_file)
                syllabus_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif syllabus_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                import docx
                doc = docx.Document(syllabus_file)
                syllabus_text = "\n".join([para.text for para in doc.paragraphs])
            elif syllabus_file.type == "text/plain":
                syllabus_text = syllabus_file.read().decode("utf-8")
        syllabus_text_input = st.text_area("Or paste syllabus text here", height=150, key="predict_syllabus_text")
        syllabus_content = syllabus_text or syllabus_text_input
        st.session_state.predict_syllabus_content = syllabus_content
        if syllabus_content:
            topics = keybert_syllabus_parser(syllabus_content)
            st.markdown(f"**Extracted Topics ({len(topics)}):**")
            st.write(", ".join(topics))
        else:
            st.info("Upload a file or paste text to see the content here.")

        # --- ML Pipeline ---
        import re
        import numpy as np
        import pandas as pd
        from sentence_transformers import SentenceTransformer
        from sklearn.metrics.pairwise import cosine_similarity
        import nltk
        nltk.download('punkt', quiet=True)
        from nltk.tokenize import sent_tokenize

        def extract_questions(text):
            # Simple regex for questions ending with ? or numbered
            questions = re.findall(r'(?:\d+\.\s*)?([A-Z][^\n\r\?]{10,}\?)', text)
            # Also split by lines and look for question-like sentences
            lines = text.split('\n')
            for line in lines:
                if len(line) > 20 and (line.strip().endswith('?') or line.strip().lower().startswith(('what', 'why', 'how', 'explain', 'describe', 'define', 'list', 'discuss', 'compare', 'differentiate'))):
                    questions.append(line.strip())
            # Remove duplicates
            questions = list(set([q.strip() for q in questions if len(q.strip()) > 10]))
            return questions

        def bloom_classifier(question):
            # Simple keyword-based classifier
            q = question.lower()
            if any(word in q for word in ["define", "list", "name", "identify", "recall"]):
                return "Remember"
            if any(word in q for word in ["explain", "summarize", "describe", "classify", "discuss"]):
                return "Understand"
            if any(word in q for word in ["apply", "solve", "use", "demonstrate", "calculate"]):
                return "Apply"
            if any(word in q for word in ["analyze", "compare", "contrast", "differentiate", "examine"]):
                return "Analyze"
            if any(word in q for word in ["evaluate", "justify", "critique", "assess", "argue"]):
                return "Evaluate"
            if any(word in q for word in ["create", "design", "formulate", "compose", "construct"]):
                return "Create"
            return "Other"

        # Engine selector for Predictive Analysis
        engine = st.radio(
            "Question Generation Engine",
            ["Extract from Previous Paper", "GPT-4 (OpenAI API)", "Local LLM (offline)"],
            index=0 if prev_content else (1 if st.secrets.get("OPENAI_API_KEY") else 2)
        )

        if prev_content and syllabus_content:
            if st.button("🔬 Analyze", type="primary"):
                with st.spinner("Running ML pipeline..."):
                    # 1. Get questions
                    if engine == "Extract from Previous Paper":
                        questions = extract_questions(prev_content)
                    else:
                        topics = keybert_syllabus_parser(syllabus_content)
                        prompt = f"Generate 20 questions for the syllabus topics: {', '.join(topics)}. Format as a list."
                        if engine == "GPT-4 (OpenAI API)" and st.secrets.get("OPENAI_API_KEY"):
                            questions_text = generate_questions_gpt4(prompt, st.secrets["OPENAI_API_KEY"])
                        else:
                            questions_text = generate_questions_local(prompt)
                        questions = [q.strip() for q in questions_text.split('\n') if len(q.strip()) > 10]
                    # 2. Extract topics
                    topics = keybert_syllabus_parser(syllabus_content)
                    # 3. Embed questions and topics
                    model = SentenceTransformer('all-MiniLM-L6-v2')
                    q_embeds = model.encode(questions)
                    t_embeds = model.encode(topics)
                    # 4. Map questions to topics
                    sim_matrix = cosine_similarity(q_embeds, t_embeds)
                    best_topic_idx = np.argmax(sim_matrix, axis=1)
                    best_topic_score = np.max(sim_matrix, axis=1)
                    mapped_topics = [topics[idx] if topics else "-" for idx in best_topic_idx]
                    # 5. Frequency analysis (count duplicate questions)
                    freq = pd.Series(questions).value_counts().to_dict()
                    # 6. Bloom's classifier
                    blooms = [bloom_classifier(q) for q in questions]
                    # 7. Display results
                    df = pd.DataFrame({
                        "Question": questions,
                        "Topic": mapped_topics,
                        "Topic Similarity": best_topic_score,
                        "Frequency": [freq[q] for q in questions],
                        "Bloom Level": blooms
                    })
                    df = df.sort_values(["Frequency", "Topic Similarity"], ascending=[False, False])
                    st.markdown("### 🧠 Predicted Most Probable Questions")
                    st.dataframe(df, use_container_width=True)
        else:
            st.warning("Please upload both previous paper and syllabus to enable analysis.")
        return

    if st.session_state.active_dashboard == "Upload Previous Paper":
        st.markdown("#### 📄 Upload Previous Question Paper")
        uploaded_file = st.file_uploader("Upload a previous paper (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"], key="prev_paper_upload")
        prev_paper_text = ""
        if uploaded_file is not None:
            if uploaded_file.type == "application/pdf":
                import PyPDF2
                reader = PyPDF2.PdfReader(uploaded_file)
                prev_paper_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                import docx
                doc = docx.Document(uploaded_file)
                prev_paper_text = "\n".join([para.text for para in doc.paragraphs])
            elif uploaded_file.type == "text/plain":
                prev_paper_text = uploaded_file.read().decode("utf-8")
        prev_paper_text_input = st.text_area("Or paste previous paper text here", height=200, key="prev_paper_text")
        content = prev_paper_text or prev_paper_text_input
        if content:
            st.markdown("**Extracted Content:**")
            st.write(content)
        else:
            st.info("Upload a file or paste text to see the content here.")

    elif st.session_state.active_dashboard == "Upload Syllabus":
        st.markdown("#### 📚 Upload Syllabus")
        uploaded_file = st.file_uploader("Upload a syllabus file (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"], key="syllabus_upload")
        syllabus_text = ""
        if uploaded_file is not None:
            if uploaded_file.type == "application/pdf":
                import PyPDF2
                reader = PyPDF2.PdfReader(uploaded_file)
                syllabus_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                import docx
                doc = docx.Document(uploaded_file)
                syllabus_text = "\n".join([para.text for para in doc.paragraphs])
            elif uploaded_file.type == "text/plain":
                syllabus_text = uploaded_file.read().decode("utf-8")
        syllabus_text_input = st.text_area("Or paste syllabus text here", height=200, key="syllabus_text")
        content = syllabus_text or syllabus_text_input
        if content:
            topics = keybert_syllabus_parser(content)
            st.markdown(f"**Extracted Topics ({len(topics)}):**")
            st.write(", ".join(topics))
        else:
            st.info("Upload a file or paste text to extract topics.")

    elif st.session_state.active_dashboard == "Auto Question Paper":
        st.markdown("#### 🚀 Auto Question Paper Generation")
        syllabus_text = st.text_area("Paste your syllabus, notes, or any text content here.", height=200, key="syllabus_input")
        col1, col2 = st.columns(2)
        with col1:
            subject_name = st.text_input("📚 Subject Name", placeholder="e.g., Database Management Systems")
            num_questions = st.slider("📝 Number of Questions", 5, 50, 20)
        with col2:
            question_types = st.multiselect(
                "🎯 Question Types",
                ["MCQ", "Short Answer", "Long Answer", "Case Study"],
                default=["MCQ", "Short Answer"]
            )
            # Engine selector
            engine = st.radio(
                "Question Generation Engine",
                ["GPT-4 (OpenAI API)", "Local LLM (offline)"],
                index=0 if st.secrets.get("OPENAI_API_KEY") else 1
            )
        if st.button("🤖 Generate Questions", type="primary", use_container_width=True):
            if syllabus_text and subject_name and question_types:
                topics = keybert_syllabus_parser(syllabus_text)
                prompt = f"Generate {num_questions} {', '.join(question_types)} questions for the subject '{subject_name}' covering these topics: {', '.join(topics)}. Format as a list."
                if engine == "GPT-4 (OpenAI API)" and st.secrets.get("OPENAI_API_KEY"):
                    with st.spinner("Generating questions using GPT-4 API..."):
                        questions_text = generate_questions_gpt4(prompt, st.secrets["OPENAI_API_KEY"])
                else:
                    with st.spinner("Generating questions using Local LLM (Mistral-7B-Instruct)..."):
                        questions_text = generate_questions_local(prompt)
                st.session_state.last_generated_questions = [q.strip() for q in questions_text.split('\n') if len(q.strip()) > 10]
                st.success(f"✅ Generated {len(st.session_state.last_generated_questions)} questions!")
            else:
                st.error("❌ Please provide syllabus content, a subject name, and select question types.")
        if 'last_generated_questions' in st.session_state:
            st.markdown("---")
            st.markdown("### 📋 Generated Questions")
            questions = st.session_state.last_generated_questions
            for i, q in enumerate(questions, 1):
                st.markdown(f"**Q{i}:** {q}")
    elif st.session_state.active_dashboard == "Collaboration":
        collaboration_dashboard()
    elif st.session_state.active_dashboard == "Admin":
        admin_dashboard()
    elif st.session_state.active_dashboard == "Super Admin":
        super_admin_dashboard()

# --- Persistent Login Integration (streamlit_app.py) ---
import streamlit_authenticator as stauth

# Example static user list (replace with DB integration later)
users = {
    "usernames": {
        "admin": {
            "name": "Admin User",
            "password": stauth.Hasher(['adminpassword']).generate()[0],
            "role": "admin"
        },
        "user": {
            "name": "Normal User",
            "password": stauth.Hasher(['userpassword']).generate()[0],
            "role": "user"
        }
    }
}

authenticator = stauth.Authenticate(
    users["usernames"],
    "questvibe_cookie", "questvibe_signature_key", cookie_expiry_days=30
)

def login_page():
    """Displays the login page using streamlit-authenticator (streamlit_app.py)"""
    name, authentication_status, username = authenticator.login("Login", "main")
    if authentication_status:
        st.session_state.current_user = {
            'id': username,
            'name': name,
            'institution': 'N/A',
            'role': users['usernames'][username]['role']
        }
        st.success(f"Welcome {name}!")
        if st.button("Logout"):
            authenticator.logout("main")
            st.session_state.current_user = None
            st.rerun()
    elif authentication_status is False:
        st.error("Username/password is incorrect")
    elif authentication_status is None:
        st.warning("Please enter your username and password")

def main():
    """Main application function"""
    st.markdown("""
    <div style="text-align: center; padding: 1rem 0;">
        <h1 style="margin-bottom: 0.5rem;">🚀 QuestVibe</h1>
        <p style="font-size: 1.2rem; color: #ffffff;">
            AI-Powered Question Paper Generation System
        </p>
    </div>
    """, unsafe_allow_html=True)

    if 'current_user' not in st.session_state:
        st.session_state.current_user = None

    if st.session_state.current_user is None:
        login_page()
    else:
        main_dashboard()

def extract_topics_from_content(content: str) -> List[str]:
    """Extract topics from syllabus content using AI processing"""
    # Simple topic extraction (can be enhanced with NLP)
    lines = content.split('\n')
    topics = []
    
    for line in lines:
        line = line.strip()
        if line:
            # Look for topic patterns
            if any(keyword in line.lower() for keyword in ['topic', 'unit', 'chapter', 'section', 'module']):
                # Extract topic name
                topic = line.replace('Topic:', '').replace('Unit:', '').replace('Chapter:', '').replace('Section:', '').replace('Module:', '').strip()
                if topic and len(topic) > 3:
                    topics.append(topic)
            elif line.startswith(('•', '-', '*', '1.', '2.', '3.', '4.', '5.')):
                # Extract bullet points as topics
                topic = line.replace('•', '').replace('-', '').replace('*', '').strip()
                if topic and len(topic) > 3:
                    topics.append(topic)
            elif len(line) > 10 and len(line) < 100:  # Potential topic length
                # Check if line looks like a topic
                if not line.endswith('.') and not line.startswith(('The', 'This', 'In', 'For', 'With')):
                    topics.append(line)
    
    # Remove duplicates and clean up
    topics = list(set(topics))
    topics = [topic for topic in topics if len(topic) > 3 and len(topic) < 100]
    
    # If no topics found, create some from content
    if not topics and content:
        words = content.split()
        # Extract potential topics from content
        for i in range(0, len(words), 10):
            if i + 5 < len(words):
                topic = ' '.join(words[i:i+5])
                topics.append(topic)
    
    return topics[:20]  # Limit to 20 topics

def generate_questions_gpt4(prompt, api_key):
    import openai
    openai.api_key = api_key
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1024,
        temperature=0.7
    )
    return response.choices[0].message.content

def generate_questions_local(prompt):
    generator = pipeline("text-generation", model="mistralai/Mistral-7B-Instruct-v0.2", torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32, device=0 if torch.cuda.is_available() else -1)
    result = generator(prompt, max_length=512, do_sample=True, temperature=0.7)
    return result[0]['generated_text']

if __name__ == "__main__":
    main() 
